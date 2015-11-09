#!/usr/bin/env python

import docx
import argparse
import datetime

title_set = False
title = None
commpany = None

def get_user_input():
	#This is where we get the position being applied for and the organization being applied to from the user
	global title, commpany
	parser = argparse.ArgumentParser(description='Script to customize cover letters')
	parser.add_argument('-p', '--position', type=str, dest='position', help='Position being applied for', required=True)
	parser.add_argument('-o', '--organization', type=str, dest='organization', help='Organization of employment', required=False)
	args = parser.parse_args()
	title = args.position
	commpany = args.organization



def customize_letter():
	global title, commpany, title_set
	#This is where we open the word template for our script
	d = docx.Document('c:\\users\\mike\\desktop\\demo.docx')
	#This is where we set the position being applied for
	for paragraph in d.paragraphs:
		if 'JB_TITLE' in paragraph.text:
			for run in paragraph.runs:
				if (('JB_TITLE' in run.text) and (title_set == False)):
					run.text = title.upper()
					title_set = True
				elif (('JB_TITLE' in run.text) and (title_set == True)):
					run.text = title
						
	#This is where we set the organization we're applying to
	#First off we check to see if the organization, being applied for has been set,
	#if it has we include it into the letter otherwise, we exclude it
	if commpany == None:
		for paragraph in d.paragraphs:
			if 'WK_PLACE' in paragraph.text:
				for run in paragraph.runs:
					if 'WK_PLACE' in run.text:
							run.text = ''
	else:	
		for paragraph in d.paragraphs:
			if 'WK_PLACE' in paragraph.text:
				for run in paragraph.runs:
					if 'WK_PLACE' in run.text:
							run.text = 'at ' + commpany


	now = datetime.datetime.now()
	date = now.strftime("%S")
	file_name = 'c:\\users\\mike\\desktop\\application-' + date + '.docx'
	print "File " + file_name + " saved"
	d.save(file_name)

if __name__ == '__main__':
	get_user_input()
	customize_letter()
